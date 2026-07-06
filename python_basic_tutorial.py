"""
產生 Python 基礎教學.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 設定全域樣式
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_title(text, level=1, color=(0, 102, 204)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    return p

def add_code(code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    # 灰底
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    return p

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# ============== 封面 ==============
title = doc.add_heading('Python 基礎教學', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('從零開始學 Python 🐶⚡')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n作者：Vans 專屬教材\n').italic = True
p.add_run('法鬥超人 整理').italic = True

doc.add_page_break()

# ============== 目錄 ==============
add_title('目錄', 1)
toc_items = [
    '一、Python 簡介',
    '二、環境安裝',
    '三、第一支程式',
    '四、變數與資料型別',
    '五、運算子',
    '六、流程控制',
    '七、迴圈',
    '八、函式',
    '九、串列 (List)',
    '十、字典 (Dictionary)',
    '十一、檔案處理',
    '十二、例外處理',
    '十三、物件導向基礎',
    '十四、實戰小範例',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ============== 一、Python 簡介 ==============
add_title('一、Python 簡介', 1)
add_para('Python 是一種直譯式、物件導向、動態語法的高階程式語言，由 Guido van Rossum 於 1991 年發布。')
add_para('主要特色：')
add_bullet('語法簡潔易讀，接近自然語言')
add_bullet('跨平台：Windows / macOS / Linux 都能跑')
add_bullet('豐富的第三方套件生態系')
add_bullet('應用領域廣：網頁、資料分析、AI、自動化、遊戲…')

add_title('為什麼選 Python？', 2)
add_para('在 AI 與資料科學的時代，Python 已經是事實上的標準語言。學會它，幾乎什麼都能做。')

# ============== 二、環境安裝 ==============
add_title('二、環境安裝', 1)
add_para('步驟：')
add_bullet('到 https://www.python.org/downloads/ 下載安裝檔')
add_bullet('安裝時記得勾選「Add Python to PATH」')
add_bullet('打開終端機輸入 python --version 驗證')

add_code('python --version\n# 輸出範例：Python 3.12.x')

add_para('推薦編輯器：VS Code、PyCharm')

# ============== 三、第一支程式 ==============
add_title('三、第一支程式', 1)
add_para('Hello World 是程式設計的傳統起點：')
add_code('print("Hello, World!")\n\n# 執行結果：\n# Hello, World!')

# ============== 四、變數與資料型別 ==============
add_title('四、變數與資料型別', 1)

add_title('變數命名', 2)
add_para('Python 不需要宣告型別，直接指派即可。')
add_code('name = "Vans"\nage = 30\nheight = 175.5\nis_teacher = True')

add_title('基本資料型別', 2)
add_code('# 字串 str\nmsg = "Hello Python"\n\n# 整數 int\nnum = 42\n\n# 浮點數 float\npi = 3.14159\n\n# 布林 bool\nis_active = True\n\n# None\nresult = None')

add_title('型別檢查', 2)
add_code('print(type(msg))    # <class \'str\'>\nprint(type(num))    # <class \'int\'>\nprint(type(pi))     # <class \'float\'>')

# ============== 五、運算子 ==============
add_title('五、運算子', 1)
add_code('# 算術運算子\na, b = 10, 3\nprint(a + b)   # 13\nprint(a - b)   # 7\nprint(a * b)   # 30\nprint(a / b)   # 3.3333...\nprint(a // b)  # 3  (整除)\nprint(a % b)   # 1  (餘數)\nprint(a ** b)  # 1000 (次方)\n\n# 比較運算子\nprint(a > b)   # True\nprint(a == b)  # False\n\n# 邏輯運算子\nprint(a > 5 and b < 5)  # True\nprint(not (a < b))       # True')

# ============== 六、流程控制 ==============
add_title('六、流程控制', 1)

add_title('if 條件判斷', 2)
add_code('score = 85\n\nif score >= 90:\n    print("優等")\nelif score >= 80:\n    print("甲等")\nelif score >= 70:\n    print("乙等")\nelse:\n    print("加油")')

add_title('match-case (Python 3.10+)', 2)
add_code('command = "start"\n\nmatch command:\n    case "start":\n        print("啟動")\n    case "stop":\n        print("停止")\n    case _:\n        print("未知指令")')

# ============== 七、迴圈 ==============
add_title('七、迴圈', 1)

add_title('for 迴圈', 2)
add_code('# 跑 5 次\nfor i in range(5):\n    print(i)  # 0 1 2 3 4\n\n# 走訪清單\nfruits = ["apple", "banana", "cherry"]\nfor fruit in fruits:\n    print(fruit)')

add_title('while 迴圈', 2)
add_code('count = 0\nwhile count < 5:\n    print(count)\n    count += 1')

add_title('break / continue', 2)
add_code('for i in range(10):\n    if i == 3:\n        continue   # 跳過這次\n    if i == 7:\n        break      # 提早結束\n    print(i)')

# ============== 八、函式 ==============
add_title('八、函式', 1)
add_para('把重複的邏輯包成函式，程式會更乾淨。')

add_code('# 基本函式\ndef greet(name):\n    return f"Hello, {name}!"\n\nprint(greet("Vans"))  # Hello, Vans!\n\n# 預設參數\ndef power(base, exp=2):\n    return base ** exp\n\nprint(power(3))     # 9\nprint(power(2, 10)) # 1024\n\n# 多回傳值\ndef calc(a, b):\n    return a + b, a - b, a * b\n\nx, y, z = calc(10, 3)\nprint(x, y, z)  # 13 7 30')

add_title('Lambda 匿名函式', 2)
add_code('square = lambda x: x * x\nprint(square(5))  # 25')

# ============== 九、串列 List ==============
add_title('九、串列 (List)', 1)
add_code('nums = [1, 2, 3, 4, 5]\n\n# 常用操作\nnums.append(6)        # 尾端新增\nnums.insert(0, 0)     # 指定位置插入\nnums.remove(3)        # 移除指定值\npopped = nums.pop()   # 移除並回傳最後一個\nprint(len(nums))      # 長度\nprint(nums[0])        # 索引取值\nprint(nums[1:4])      # 切片\n\n# 串列生成式\nsquares = [x**2 for x in range(5)]\nprint(squares)  # [0, 1, 4, 9, 16]')

# ============== 十、字典 ==============
add_title('十、字典 (Dictionary)', 1)
add_code('user = {\n    "name": "Vans",\n    "age": 30,\n    "skills": ["Python", "AI", "Teaching"]\n}\n\nprint(user["name"])            # Vans\nuser["email"] = "vans@ai.com"  # 新增\n\nfor key, value in user.items():\n    print(f"{key}: {value}")')

# ============== 十一、檔案處理 ==============
add_title('十一、檔案處理', 1)
add_code('# 寫入檔案\nwith open("note.txt", "w", encoding="utf-8") as f:\n    f.write("Hello Python\\n")\n    f.write("第二行\\n")\n\n# 讀取檔案\nwith open("note.txt", "r", encoding="utf-8") as f:\n    content = f.read()\n    print(content)')

add_para('with 會自動關檔，比傳統的 open/close 安全。')

# ============== 十二、例外處理 ==============
add_title('十二、例外處理', 1)
add_code('try:\n    num = int(input("請輸入數字: "))\n    result = 100 / num\nexcept ValueError:\n    print("這不是數字！")\nexcept ZeroDivisionError:\n    print("不能除以 0")\nelse:\n    print(f"結果是 {result}")\nfinally:\n    print("執行結束")')

# ============== 十三、物件導向 ==============
add_title('十三、物件導向基礎', 1)
add_code('class Dog:\n    def __init__(self, name, breed):\n        self.name = name\n        self.breed = breed\n\n    def bark(self):\n        return f"{self.name}: 汪汪！"\n\nmy_dog = Dog("法鬥超人", "法鬥")\nprint(my_dog.bark())  # 法鬥超人: 汪汪！')

add_title('繼承', 2)
add_code('class Bulldog(Dog):\n    def __init__(self, name):\n        super().__init__(name, "Bulldog")\n\n    def bark(self):\n        return f"{self.name}: 嗶嗶！(用鼻子發聲)"')

# ============== 十四、實戰範例 ==============
add_title('十四、實戰小範例', 1)

add_title('範例 1：FizzBuzz', 2)
add_code('for i in range(1, 21):\n    if i % 15 == 0:\n        print("FizzBuzz")\n    elif i % 3 == 0:\n        print("Fizz")\n    elif i % 5 == 0:\n        print("Buzz")\n    else:\n        print(i)')

add_title('範例 2：猜數字遊戲', 2)
add_code('import random\n\nanswer = random.randint(1, 100)\nattempts = 0\n\nwhile True:\n    guess = int(input("猜 1~100: "))\n    attempts += 1\n\n    if guess < answer:\n        print("大一點")\n    elif guess > answer:\n        print("小一點")\n    else:\n        print(f"答對了！共猜了 {attempts} 次")\n        break')

add_title('範例 3：計算字數', 2)
add_code('text = "Python is awesome. Python is easy."\nwords = text.split()\ncounter = {}\n\nfor w in words:\n    counter[w] = counter.get(w, 0) + 1\n\nprint(counter)')

# ============== 結語 ==============
doc.add_page_break()
add_title('結語', 1)
add_para('Python 是一條值得投資的長線。掌握基礎後，可以再往以下方向發展：')

add_bullet('網頁開發：Django、Flask、FastAPI')
add_bullet('資料分析：Pandas、NumPy')
add_bullet('機器學習 / AI：scikit-learn、PyTorch、TensorFlow')
add_bullet('自動化腳本：Selenium、BeautifulSoup')
add_bullet('雲端與 DevOps')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n🐶 寫程式就像養法鬥，需要耐心，但回報超甜 ⚡')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(255, 102, 0)

# 儲存
output_path = 'Python_基礎教學.docx'
doc.save(output_path)
print(f'✅ 檔案產生完成：{output_path}')
